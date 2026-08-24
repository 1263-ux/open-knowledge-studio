#!/usr/bin/env python3
"""Build a small, source-traceable DOCX from a curated JSON outline.

The helper is intentionally deterministic: recall and claim selection happen
in the skill workflow, while this file only renders the approved outline.
"""
from __future__ import annotations

import argparse
import json
import os
from pathlib import Path
import tempfile

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor
from docx.oxml.ns import qn


def _set_run_font(run, name: str = "Aptos") -> None:
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")


def _add_text(paragraph, text: str, bold: bool = False) -> None:
    run = paragraph.add_run(str(text))
    run.bold = bold
    _set_run_font(run)


def _configure(document: Document) -> None:
    section = document.sections[0]
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)

    styles = document.styles
    normal = styles["Normal"]
    normal.font.name = "Aptos"
    normal._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(6)
    for style_name, size, color in (
        ("Title", 24, RGBColor(15, 23, 42)),
        ("Heading 1", 16, RGBColor(15, 23, 42)),
        ("Heading 2", 12, RGBColor(30, 64, 175)),
    ):
        style = styles[style_name]
        style.font.name = "Aptos Display"
        style._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.color.rgb = color


def build(outline: dict, output: Path) -> None:
    document = Document()
    _configure(document)

    title = document.add_paragraph(style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    _add_text(title, outline.get("title", "Knowledge brief"))

    if outline.get("subtitle"):
        subtitle = document.add_paragraph()
        subtitle.paragraph_format.space_after = Pt(14)
        _add_text(subtitle, outline["subtitle"])

    if outline.get("summary"):
        heading = document.add_paragraph(style="Heading 1")
        _add_text(heading, "Executive summary")
        summary = document.add_paragraph()
        _add_text(summary, outline["summary"])

    for section in outline.get("sections", []):
        heading = document.add_paragraph(style="Heading 1")
        _add_text(heading, section.get("heading", "Section"))
        for paragraph_text in section.get("paragraphs", []):
            paragraph = document.add_paragraph()
            _add_text(paragraph, paragraph_text)
        for bullet in section.get("bullets", []):
            paragraph = document.add_paragraph(style="List Bullet")
            _add_text(paragraph, bullet)
        table_rows = [row for row in section.get("table", []) if row]
        if table_rows:
            table = document.add_table(rows=1, cols=len(table_rows[0]))
            table.style = "Light Shading Accent 1"
            table.autofit = False
            column_width = Inches(6.9 / len(table_rows[0]))
            for cell, value in zip(table.rows[0].cells, table_rows[0]):
                cell.text = str(value)
                cell.width = column_width
            for row in table_rows[1:]:
                cells = table.add_row().cells
                for cell, value in zip(cells, row):
                    cell.text = str(value)
                    cell.width = column_width
            for row in table.rows:
                for cell in row.cells:
                    cell.width = column_width

    sources = outline.get("sources", [])
    if sources:
        heading = document.add_paragraph(style="Heading 1")
        _add_text(heading, "Sources")
        for source in sources:
            paragraph = document.add_paragraph(style="List Bullet")
            label = source.get("label", "Source")
            path = source.get("path", "")
            _add_text(paragraph, f"{label}: ")
            _add_text(paragraph, path)

    output.parent.mkdir(parents=True, exist_ok=True)
    fd, temp_name = tempfile.mkstemp(prefix=f".{output.stem}-", suffix=".docx", dir=output.parent)
    os.close(fd)
    temp_path = Path(temp_name)
    try:
        document.save(temp_path)
        os.replace(temp_path, output)
    finally:
        temp_path.unlink(missing_ok=True)


def main() -> int:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("--outline", required=True, type=Path, help="Curated JSON outline")
    parser.add_argument("--output", required=True, type=Path, help="Output .docx path")
    args = parser.parse_args()
    outline = json.loads(args.outline.read_text(encoding="utf-8"))
    if not isinstance(outline, dict):
        raise SystemExit("outline must be a JSON object")
    build(outline, args.output)
    print(args.output)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
